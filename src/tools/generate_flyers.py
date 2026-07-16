#!/usr/bin/env python3
"""Generate Word document flyers for EnglishConnect classes."""

from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

PROJECT_ROOT = Path(__file__).resolve().parent.parent.parent
QR_DIR = PROJECT_ROOT / "content" / "qr-codes"
OUTPUT_DIR = PROJECT_ROOT / "content" / "flyers"
WHATSAPP_QR = PROJECT_ROOT / "documentation" / "whatsapp-group-ec1.png"


def set_paragraph_spacing(paragraph, before=0, after=0):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)


def add_centered_text(doc, text, size=16, bold=False, color=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)
    set_paragraph_spacing(p, before=2, after=2)
    return p


def add_centered_image(doc, image_path, width_inches=2.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))
    set_paragraph_spacing(p, before=4, after=4)
    return p


def generate_ec1_flyer():
    doc = Document()

    # Title
    add_centered_text(doc, "English Connect 1B", size=28, bold=True)
    add_centered_text(doc, "Free English Classes / Clases de Ingles Gratis", size=14, color=(80, 80, 80))

    # Teachers
    add_centered_text(doc, "Hno Molen y Elder Peterson", size=14)

    # Schedule
    add_centered_text(doc, "Tuesdays & Thursdays / Martes y Jueves", size=13, bold=True)
    add_centered_text(doc, "5:00 - 6:00 PM  &  7:00 - 8:00 PM", size=13)

    # Location
    add_centered_text(doc, "7699 Fallbrook Dr, Houston, TX 77086", size=12, color=(80, 80, 80))

    # Spacer
    add_centered_text(doc, "", size=6)

    # EnglishConnect QR
    add_centered_text(doc, "Practice Online / Practica en Linea", size=14, bold=True)
    add_centered_image(doc, QR_DIR / "ec1" / "ec1.png", width_inches=2.2)
    add_centered_text(doc, "englishconnect.vozloop.com", size=10, color=(100, 100, 100))

    # Spacer
    add_centered_text(doc, "", size=6)

    # WhatsApp QR
    add_centered_text(doc, "Join Our WhatsApp Group / Unete al Grupo de WhatsApp", size=14, bold=True)
    add_centered_image(doc, WHATSAPP_QR, width_inches=2.2)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    out_path = OUTPUT_DIR / "EC1B-flyer.docx"
    doc.save(str(out_path))
    print(f"EC1 flyer saved to: {out_path}")
    return out_path


def generate_ec2_flyer():
    doc = Document()

    # Title
    add_centered_text(doc, "English Connect 2B", size=28, bold=True)
    add_centered_text(doc, "Free English Classes / Clases de Ingles Gratis", size=14, color=(80, 80, 80))

    # Schedule
    add_centered_text(doc, "Tuesdays & Thursdays / Martes y Jueves", size=13, bold=True)
    add_centered_text(doc, "5:00 - 6:00 PM  &  7:00 - 8:00 PM", size=13)

    # Location
    add_centered_text(doc, "7699 Fallbrook Dr, Houston, TX 77086", size=12, color=(80, 80, 80))

    # Spacer
    add_centered_text(doc, "", size=8)

    # EnglishConnect QR
    add_centered_text(doc, "Practice Online / Practica en Linea", size=14, bold=True)
    add_centered_image(doc, QR_DIR / "ec2" / "ec2.png", width_inches=2.8)
    add_centered_text(doc, "englishconnect.vozloop.com", size=10, color=(100, 100, 100))

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    out_path = OUTPUT_DIR / "EC2B-flyer.docx"
    doc.save(str(out_path))
    print(f"EC2 flyer saved to: {out_path}")
    return out_path


if __name__ == "__main__":
    generate_ec1_flyer()
    generate_ec2_flyer()
